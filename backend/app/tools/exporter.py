from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import tempfile


def export_docx(report_text):

    temp_file = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx"
    )

    doc = Document()

    for line in report_text.split("\n"):

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)

        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)

        elif line.strip():
            doc.add_paragraph(line)

    doc.save(temp_file.name)

    return temp_file.name


def export_pdf(report_text):

    temp_file = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".pdf"
    )

    doc = SimpleDocTemplate(temp_file.name)

    styles = getSampleStyleSheet()

    content = []

    for line in report_text.split("\n"):

        if line.strip():

            content.append(
                Paragraph(line, styles["BodyText"])
            )

            content.append(
                Spacer(1, 5)
            )

    doc.build(content)

    return temp_file.name